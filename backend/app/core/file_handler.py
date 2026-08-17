"""
通用文件处理工具

提供文件上传校验、文本读取、临时保存、文档内容提取等公共能力，
供知识库上传、接口导入、数据导入等模块复用。
"""
import os
import tempfile
import logging
from typing import Optional

from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)


class FileHandler:
    """通用文件处理"""

    # 允许的文本类扩展名
    TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml", ".xml", ".html"}
    # 允许的文档类扩展名
    DOC_EXTENSIONS = {".docx", ".pdf", ".doc"}
    # 允许的导入类扩展名
    IMPORT_EXTENSIONS = {".json", ".har", ".jmx", ".yaml", ".yml", ".postman_collection"}
    # 全部允许的扩展名
    ALLOWED_EXTENSIONS = TEXT_EXTENSIONS | DOC_EXTENSIONS | IMPORT_EXTENSIONS
    # 最大文件大小 10MB
    MAX_SIZE = 10 * 1024 * 1024

    @staticmethod
    def validate(file: UploadFile, allowed_extensions: Optional[set] = None, max_size: int = None) -> None:
        """
        校验文件扩展名和大小。

        Args:
            file: UploadFile 对象
            allowed_extensions: 允许的扩展名集合（含点号，小写），为 None 时使用 ALLOWED_EXTENSIONS
            max_size: 最大字节数，为 None 时使用 MAX_SIZE

        Raises:
            HTTPException: 校验失败时抛出 400
        """
        allowed = allowed_extensions or FileHandler.ALLOWED_EXTENSIONS
        max_bytes = max_size or FileHandler.MAX_SIZE

        filename = file.filename or ""
        ext = os.path.splitext(filename)[1].lower()

        if ext not in allowed:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型 {ext}，允许的类型：{', '.join(sorted(allowed))}",
            )

        # 检查文件大小（通过 seek 到末尾获取大小）
        file.file.seek(0, 2)
        size = file.file.tell()
        file.file.seek(0)

        if size > max_bytes:
            raise HTTPException(
                status_code=400,
                detail=f"文件大小 {size} 字节超过限制 {max_bytes} 字节",
            )

    @staticmethod
    def read_text(file: UploadFile, encoding: str = "utf-8") -> str:
        """
        读取文本文件内容。

        Args:
            file: UploadFile 对象
            encoding: 文件编码

        Returns:
            文件文本内容
        """
        content = file.file.read()
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            # 尝试其他常见编码
            for enc in ("gbk", "gb2312", "latin-1"):
                try:
                    return content.decode(enc)
                except UnicodeDecodeError:
                    continue
            return content.decode(encoding, errors="ignore")

    @staticmethod
    def save_to_temp(file: UploadFile, suffix: Optional[str] = None) -> str:
        """
        将上传文件保存到临时文件，返回临时文件路径。

        Args:
            file: UploadFile 对象
            suffix: 临时文件后缀，为 None 时使用原文件扩展名

        Returns:
            临时文件绝对路径
        """
        if suffix is None:
            suffix = os.path.splitext(file.filename or "")[1]

        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        try:
            with os.fdopen(fd, "wb") as f:
                f.write(file.file.read())
        except Exception:
            os.unlink(tmp_path)
            raise
        finally:
            file.file.seek(0)

        return tmp_path

    @staticmethod
    def extract_content(file_path: str) -> str:
        """
        从文档文件中提取文本内容，支持 .txt/.md/.docx/.pdf。

        Args:
            file_path: 文件绝对路径

        Returns:
            提取的文本内容
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext in (".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml"):
            return FileHandler._read_text_file(file_path)
        elif ext == ".docx":
            return FileHandler._extract_docx(file_path)
        elif ext == ".pdf":
            return FileHandler._extract_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件类型：{ext}")

    @staticmethod
    def _read_text_file(file_path: str) -> str:
        """读取文本文件，自动尝试多种编码"""
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """从 .docx 文件提取文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx 未安装，无法提取 .docx 文件")
            raise HTTPException(status_code=400, detail="服务器未安装 python-docx，无法解析 Word 文档")
        except Exception as e:
            logger.error(f"提取 docx 内容失败: {e}")
            raise HTTPException(status_code=400, detail=f"Word 文档解析失败：{e}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """从 .pdf 文件提取文本"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n".join(texts)
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                texts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
                return "\n".join(texts)
            except ImportError:
                logger.warning("pypdf/PyPDF2 未安装，无法提取 .pdf 文件")
                raise HTTPException(status_code=400, detail="服务器未安装 pypdf，无法解析 PDF 文档")
        except Exception as e:
            logger.error(f"提取 pdf 内容失败: {e}")
            raise HTTPException(status_code=400, detail=f"PDF 文档解析失败：{e}")

    @staticmethod
    def cleanup_temp(file_path: str) -> None:
        """清理临时文件"""
        try:
            if file_path and os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            logger.warning(f"清理临时文件失败 {file_path}: {e}")
